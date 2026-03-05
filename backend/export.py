import json
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import pysrt


def export_txt(transcript: str, summary: str, filename: str, content: str = "both") -> bytes:
    """Export transcript and/or summary as plain text."""
    parts = []
    
    # Build text output depending on requested content
    if content == "transcript":
        parts.append(f"Transcript: {filename}\n")
        if transcript:
            parts.append(f"FULL TRANSCRIPT:\n{transcript}")
    elif content == "summary":
        parts.append(f"Summary: {filename}\n")
        if summary:
            parts.append(f"SUMMARY:\n{summary}")
    else:  # both
        parts.append(f"Transcript & Summary: {filename}\n")
        if summary:
            parts.append(f"SUMMARY:\n{summary}\n")
        if transcript:
            parts.append(f"FULL TRANSCRIPT:\n{transcript}")
    
    # Return encoded text file content
    return "\n".join(parts).encode('utf-8')


def export_pdf(transcript: str, summary: str, filename: str, content: str = "both") -> bytes:
    """Export transcript and/or summary as a PDF document."""
    buffer = BytesIO()
    
    # Create a PDF document in memory
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Build document structure using reportlab flowables
    if content == "transcript":
        story.append(Paragraph(f"Transcript: {filename}", styles['Title']))
        story.append(Spacer(1, 12))
        if transcript:
            story.append(Paragraph("FULL TRANSCRIPT:", styles['Heading2']))
            story.append(Paragraph(transcript, styles['Normal']))
    elif content == "summary":
        story.append(Paragraph(f"Summary: {filename}", styles['Title']))
        story.append(Spacer(1, 12))
        if summary:
            story.append(Paragraph("SUMMARY:", styles['Heading2']))
            story.append(Paragraph(summary, styles['Normal']))
    else:  # both
        story.append(Paragraph(f"Transcript & Summary: {filename}", styles['Title']))
        story.append(Spacer(1, 12))
        if summary:
            story.append(Paragraph("SUMMARY:", styles['Heading2']))
            story.append(Paragraph(summary, styles['Normal']))
            story.append(Spacer(1, 12))
        if transcript:
            story.append(Paragraph("FULL TRANSCRIPT:", styles['Heading2']))
            story.append(Paragraph(transcript, styles['Normal']))
    
    # Build PDF and return file bytes
    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def export_docx(transcript: str, summary: str, filename: str, content: str = "both") -> bytes:
    """Export transcript and/or summary as a DOCX file."""
    doc = Document()
    
    # Populate document content based on requested section
    if content == "transcript":
        doc.add_heading(f"Transcript: {filename}", 0)
        if transcript:
            doc.add_heading("Full Transcript", level=1)
            doc.add_paragraph(transcript)
    elif content == "summary":
        doc.add_heading(f"Summary: {filename}", 0)
        if summary:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(summary)
    else:  # both
        doc.add_heading(f"Transcript & Summary: {filename}", 0)
        if summary:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(summary)
        if transcript:
            doc.add_heading("Full Transcript", level=1)
            doc.add_paragraph(transcript)
    
    # Save document into an in-memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def export_srt(word_timestamps: str, transcript: str) -> bytes:
    """Export transcript as SRT subtitle format using word-level timestamps."""
    try:
        # Load word timestamps produced by the transcription engine
        words = json.loads(word_timestamps)
        subs = pysrt.SubRipFile()
        
        # Group words into subtitle chunks (every ~10 words)
        chunk_size = 10
        for i in range(0, len(words), chunk_size):
            chunk = words[i:i + chunk_size]
            if not chunk:
                continue
            
            # Convert seconds to milliseconds for SRT format
            start_ms = int(chunk[0]["start"] * 1000)
            end_ms = int(chunk[-1]["end"] * 1000)
            
            # Combine words into subtitle text
            text = " ".join([w["word"] for w in chunk])
            
            sub = pysrt.SubRipItem(
                index=len(subs) + 1,
                start=pysrt.SubRipTime(milliseconds=start_ms),
                end=pysrt.SubRipTime(milliseconds=end_ms),
                text=text
            )
            subs.append(sub)
        
        return str(subs).encode('utf-8')
    except:
        # Fallback if timestamps are missing or invalid
        return f"1\n00:00:00,000 --> 00:00:10,000\n{transcript[:100]}".encode('utf-8')